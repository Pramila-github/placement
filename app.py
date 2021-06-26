import streamlit as st
import numpy as np
import pandas as pd
import pickle
import os
import base64
import docx
from PIL import Image

dataset = pd.read_csv('placementdata.csv')
def convert_to_int(word):
    word_dict = {'Good':1, 'Bad':0,'Yes':1,'No':0,'completed':1,'none':0,'Placed':1,'Not PLaced':0}
    return word_dict[word]


dataset = dataset.drop('sl_no', axis=1)

# catgorising col for further labelling
dataset["gender"] = dataset["gender"].astype('category')
dataset["comSkill"] = dataset["comSkill"].astype('category')
dataset["ssc_b"] = dataset["ssc_b"].astype('category')
dataset["hsc_b"] = dataset["hsc_b"].astype('category')
dataset["degree_t"] = dataset["degree_t"].astype('category')
dataset["Internship"] = dataset["Internship"].astype('category')
#dataset["status"] = dataset["status"].astype('category')
dataset["sports"] = dataset["sports"].astype('category')
dataset["hsc_s"] = dataset["hsc_s"].astype('category')

# labelling the columns
dataset["gender"] = dataset["gender"].cat.codes
dataset["comSkill"] = dataset["comSkill"].cat.codes
dataset["ssc_b"] = dataset["ssc_b"].cat.codes
dataset["hsc_b"] = dataset["hsc_b"].cat.codes
dataset["degree_t"] = dataset["degree_t"].cat.codes
dataset["Internship"] = dataset["Internship"].cat.codes
dataset["sports"] = dataset["sports"].cat.codes
dataset["hsc_s"] = dataset["hsc_s"].cat.codes
#dataset["status"] = dataset["status"].cat.codes


# selecting the features and labels
X = dataset.iloc[:, :-1].values
Y=dataset.loc[:,['salary']]
from sklearn.metrics import accuracy_score

from sklearn.model_selection import train_test_split
X_train, X_test, Y_train, Y_test = train_test_split(X, Y, test_size=0.2)



from sklearn.ensemble import RandomForestClassifier
rand_forest = RandomForestClassifier()
rand_forest=rand_forest.fit(X_train,Y_train)
print(rand_forest.predict([[0,1,100.00,0,100.33,1,0,100.4,1,6,0,86,66.28,0]]))

# Saving model to disk
pickle.dump(rand_forest, open('final_model.pkl','wb'))
# Loading model to compare the results
#print(model.predict([[0,convert_to_int('Good'),39,0,58,0,0,58,0,4,convert_to_int('none'),58,58,convert_to_int('No')]]))
st.beta_set_page_config(initial_sidebar_state="expanded")
#st.markdown(f'''<style>.sidebar .sidebar-content {{width: 375px;}}</style>''',unsafe_allow_html=True)
models = pickle.load(open('final_model.pkl', 'rb'))

def home():
    return "welcome"


def predict(gender,comSkill1, ssc_p,ssc_b, hsc_p,hsc_b,hsc_s, degree_p,degree_t, No_certi, internships1,etest_p, placetest_p,sports1):
    prediction = models.predict([[gender,comSkill1, ssc_p,ssc_b, hsc_p,hsc_b,hsc_s, degree_p,degree_t, No_certi, internships1,etest_p, placetest_p,sports1]])
    print(prediction)
    return prediction


def convert_to_int(word):
    word_dict = {'Good': 1, 'Bad': 0, 'yes': 1, 'no': 0, 'completed': 1, 'none': 0,'Male':0,'Female':1,'Central':0, 'Other':1,'Science':0, 'Arts':1, 'biology':1, 'Others':0,'Computer science Engineering':0, ' B.Tech IT':1, 'Electronics':2}
    return word_dict[word]

def get_binary_file_downloader_html(bin_file, file_label='File'):
        with open(bin_file, 'rb') as f:
            data = f.read()           
        bin_str = base64.b64encode(data).decode()
        href = f'<a href="data:application/octet-stream;base64,{bin_str}" download="{os.path.basename(bin_file)}">Download {file_label}</a>'
        return href


def main():
    st.title("STUDENT PLACEMENT PREDICTION⭐")
    st.sidebar.header('STUDENT DETAILS💻')
    name = st.sidebar.text_input('Name 📋')
    sl_no = st.sidebar.text_input('Register Number 📋')
    gender = st.sidebar.radio("Gender 👫", ('Male', 'Female'))
    gender1=convert_to_int(gender)
    comSkill = st.sidebar.radio("Communication Skills 💬🎤", ('Good', 'Bad'))
    comSkill1 = convert_to_int(comSkill)
    ssc_p = st.sidebar.slider('SSC percentage', 0, 1, 100)
    ssc_b = st.sidebar.radio("SSC Board 📚", ('Central', 'Other'))
    ssc_b1 = convert_to_int(ssc_b)
    hsc_p = st.sidebar.slider('HSC percentage', 0, 1, 100)
    hsc_b = st.sidebar.radio("HSC Board 📚", ('Central', 'Other'))
    hsc_b1 = convert_to_int(hsc_b)
    hsc_s = st.sidebar.radio("HSC Group", ('Science', 'Arts', 'Biology', 'Other'))
    hsc_s = convert_to_int(hsc_s)
    degree_t = st.radio("Degree(Engineering) 🎓", ('Computer science Engineering', ' B.Tech IT', 'Electronics','Other'))
    degree_t1 =convert_to_int(degree_t)
    degree_p = st.slider('Degree percentage', 0, 1, 100)
    No_certi = st.number_input('Number of certifications 🏆📜', min_value=0, max_value=100, value=0, step=1)
    internships = st.radio("Internships 👨‍💻", ('completed', 'none'))
    internships1 = convert_to_int(internships)
    etest_p = st.slider('E-tests(aptitude)  🏁', 0, 1, 100)
    sports = st.radio("Through Sports ⚽️🏏🏃‍♂️ ", ('yes', 'no'))
    sports1 = convert_to_int(sports)
    placetest_p = st.slider('Placement Percentage 🎓🎓 ', 0, 1, 100)
    st.radio("Interest about placement ✔️", ('Interested', 'Not Interested'))

    result = ""
    if st.button("Submit"):
        result = predict(gender1,comSkill1, ssc_p,ssc_b1, hsc_p,hsc_b1,hsc_s, degree_p,degree_t1,No_certi, internships1,etest_p, placetest_p,sports1)
        print(result)
    

        st.subheader("Entered Details")
        st.write("Name 📋 : ", name)
        st.write("Register Number 📋 :", sl_no)
        st.write("SSC percentage 📚 :", ssc_p)
        st.write("HSC percentage 📚 :", hsc_p)
        st.write("Degree percentage 🎓 :", degree_p)
        st.write("E-Test(Aptitude) 🏁 :", etest_p)
        st.write("Placement Percentage 🎓  :", placetest_p)
        st.write("Communication Skills 💬🎤 :", comSkill)
        st.write("Internship status 👨‍💻 :", internships)
        st.write("Number of certifications 🏆📜 :", No_certi)
        st.write("Through sports ⚽️🏏🏃‍♂️  :", sports)

        if(result>[0]):
          st.success('Predicted Results :- You will be Placed with {} per annum'.format("%.1f" % result))
        else:
           st.success('Predicted Results :- Not Placed') 
        
    if(st.checkbox("Generate Resume")): 
        with st.beta_expander("1. Image Upload"):
            uploaded_file = st.file_uploader(label="Upload your profile picture",type=['png','jpeg','jpg'])
            if uploaded_file is None:
                uploaded_file='user.png'
        with st.beta_expander("2. Personal Information"):
            col1,col2,col3 = st.beta_columns(3) 
            with col1:
                prof = st.text_input("Profession")
            with col2:
                city= st.text_input("City")
            with col3:
                state= st.text_input("State")
            col4,col5,col6 = st.beta_columns(3) 
            with col4:
                pname = st.text_input("Parent/Guardian Name")
            with col5:
                dob= st.text_input("Date Of Birth")
            with col6:
                lang= st.text_input("Known Languages")
            col7,col8,col9 = st.beta_columns(3)
            with col7:
                email = st.text_input("Email Id")
            with col8:
                contact= st.text_input("Contact Number")
            with col9:
                ex= st.text_input("Extra-curricular Activities")
                
        with st.beta_expander("3. Academic Information"):
            col1,col2,col3 = st.beta_columns(3) 
            with col1:
                ssc= st.text_input("SSC completion year")
            with col2:
                hsc= st.text_input("HSC completion year")
            with col3:
                deg= st.text_input("Degree Period (Ex:2020-2024)")
            college=st.text_input("College/University Name")
            intern=st.text_area("Details about Internships/Certifications")
            proj=st.text_area("Details about Academic Projects")
            skills=st.text_area("Areas of Strength/Skills")
        st.write("")    
        if st.button("Finish"):
            doc = docx.Document('Template.docx')
            from docx.shared import Pt        
            style = doc.styles['Normal']
            font = style.font
            font.name = 'Times New Roman'
            font.size = Pt(12)
            tables = doc.tables
            img = tables[0].rows[0].cells[0].add_paragraph()
            r = img.add_run()
            r.add_picture(uploaded_file, width=docx.shared.Inches(2), height=docx.shared.Inches(2))
            for paragraph in doc.paragraphs:
                if 'NAME' in paragraph.text: 
                    paragraph.text=""
                    paragraph.add_run(name).bold=True
                if 'PROFESSION' in paragraph.text: 
                    paragraph.text=""
                    paragraph.add_run(prof).bold=True
                if 'City' in paragraph.text: 
                    paragraph.text=city
                if 'State' in paragraph.text: 
                    paragraph.text=state
                if 'Contact No: ' in paragraph.text: 
                    paragraph.add_run(contact)
                if 'Email: ' in paragraph.text: 
                    paragraph.add_run(email)
                if 'Degree' in paragraph.text: 
                    paragraph.text = paragraph.text.replace("Degree", degree_t)
                if 'college' in paragraph.text: 
                    paragraph.text = paragraph.text.replace("college", college)
                if '0000-0000 ' in paragraph.text: 
                    paragraph.text=paragraph.text.replace("0000-0000 ", deg)
                if 'deg_p' in paragraph.text: 
                    paragraph.text=paragraph.text.replace("deg_p", str(degree_p))    
                if 'hsc_b' in paragraph.text: 
                    paragraph.text=paragraph.text.replace("hsc_b",hsc_b)   
                if '1212' in paragraph.text: 
                    paragraph.text=paragraph.text.replace("1212",hsc)
                if 'hsc_p' in paragraph.text: 
                    paragraph.text=paragraph.text.replace("hsc_p", str(hsc_p)) 
                if 'ssc_b' in paragraph.text: 
                    paragraph.text=paragraph.text.replace("ssc_b",ssc_b)   
                if '1010' in paragraph.text: 
                    paragraph.text=paragraph.text.replace("1010",ssc)
                if 'ssc_p' in paragraph.text: 
                    paragraph.text=paragraph.text.replace("ssc_p", str(ssc_p))
                if 'certificate' in paragraph.text: 
                    paragraph.text=paragraph.text.replace("certificate",intern)
                if 'project' in paragraph.text: 
                    paragraph.text=paragraph.text.replace("project",proj)
                if 'skill' in paragraph.text: 
                    paragraph.text=paragraph.text.replace("skill",skills)
                if 'pname' in paragraph.text: 
                    paragraph.text=paragraph.text.replace("pname",pname)
                if 'dob' in paragraph.text: 
                    paragraph.text=paragraph.text.replace("dob",dob)
                if 'lang' in paragraph.text: 
                    paragraph.text=paragraph.text.replace("lang",lang)
                if 'ext' in paragraph.text: 
                    paragraph.text=paragraph.text.replace("ext",ex)
            
            doc.save('Resume.docx') 
            st.markdown(get_binary_file_downloader_html('Resume.docx', 'Your Resume'), unsafe_allow_html=True)  

    
if __name__ == "__main__":
    main()
